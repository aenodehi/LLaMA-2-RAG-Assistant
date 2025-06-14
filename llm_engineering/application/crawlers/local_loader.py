import os
import uuid
from pathlib import Path
from typing import List, Optional, Dict
from datetime import datetime
from llm_engineering.domain.cleaned_documents import LocalFileDocument

# Text extraction functions (keep your existing ones)
def load_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz  # PyMuPDF
    try:
        with fitz.open(filepath) as doc:
            return "".join(page.get_text() for page in doc)
    except Exception as e:
        raise ValueError(f"Failed to load PDF {filepath}: {str(e)}")

def load_text_from_txt(filepath: str) -> str:
    """Read text from plain text file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(filepath, 'r', encoding='latin-1') as f:
            return f.read()

def load_text_from_docx(filepath: str) -> str:
    """Extract text from Word document."""
    from docx import Document
    try:
        doc = Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs if para.text)
    except Exception as e:
        raise ValueError(f"Failed to load DOCX {filepath}: {str(e)}")

# New text extraction functions
def load_text_from_html(filepath: str) -> str:
    """Extract clean text from HTML file."""
    from bs4 import BeautifulSoup
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            return soup.get_text(separator='\n', strip=True)
    except Exception as e:
        raise ValueError(f"Failed to load HTML {filepath}: {str(e)}")

def load_text_from_image(filepath: str) -> str:
    """Extract text from images using OCR."""
    try:
        from PIL import Image
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
        return pytesseract.image_to_string(Image.open(filepath))
    except Exception as e:
        raise ValueError(f"Failed to process image {filepath}: {str(e)}")

class LocalFileLoader:
    # Supported file extensions and their handlers
    FILE_HANDLERS = {
        '.pdf': load_text_from_pdf,
        '.txt': load_text_from_txt,
        '.docx': load_text_from_docx,
        '.html': load_text_from_html,
        '.htm': load_text_from_html,
        '.jpg': load_text_from_image,
        '.jpeg': load_text_from_image,
        '.png': load_text_from_image,
        '.md': load_text_from_txt
    }

    def __init__(self, folder_path: str, enable_ocr: bool = True):
        """
        Initialize the file loader.
        
        Args:
            folder_path: Path to the folder containing documents
            enable_ocr: Whether to enable OCR for image files
        """
        self.folder_path = Path(folder_path).expanduser().absolute()
        self.enable_ocr = enable_ocr

    def _get_file_metadata(self, filepath: Path) -> Dict:
        """Extract common file metadata."""
        stat = filepath.stat()
        return {
            'file_size': stat.st_size,
            'last_modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'file_type': filepath.suffix[1:].upper()
        }

    def load_documents(self) -> List[LocalFileDocument]:
        """Load and process all supported documents in the folder."""
        if not self.folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {self.folder_path}")

        documents = []
        for filepath in self.folder_path.iterdir():
            if not filepath.is_file():
                continue

            suffix = filepath.suffix.lower()
            if suffix not in self.FILE_HANDLERS:
                continue
            if not self.enable_ocr and suffix in {'.jpg', '.jpeg', '.png'}:
                continue

            print(f"[SCAN] Found file: {filepath.name}")

            try:
                handler = self.FILE_HANDLERS[suffix]
                content = handler(str(filepath))
                
                print(f"[DEBUG] Processed {filepath.name}")
                print(f"[DEBUG] Extracted content (first 300 chars): {content[:300]}")


                if not content.strip():
                    continue

                metadata = self._get_file_metadata(filepath)
                documents.append(LocalFileDocument(
                    id=str(uuid.uuid4()),
                    content=content,
                    file_path=str(filepath),
                    **metadata
                ))

            except Exception as e:
                print(f"Skipping {filepath.name} due to error: {str(e)}")
                continue

        return documents
